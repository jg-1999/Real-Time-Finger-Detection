from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("assets") / "Real-Time_Finger_Detection_Report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths_in):
            if index < len(row.cells):
                set_cell_width(row.cells[index], width)
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Informe tecnico - Real-Time Finger Detection")
    set_run_font(run, size=9, color="666666")


def add_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.167
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.167
        paragraph.add_run(item)


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9, color="333333")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        set_cell_shading(header_cells[index], "F2F4F7")
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                set_run_font(run, size=10, color="0B2545")

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.10
                for run in paragraph.runs:
                    set_run_font(run, size=10, color="222222")

    set_table_geometry(table, widths)
    document.add_paragraph()


def add_callout(document: Document, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title + ": ")
    run.bold = True
    set_run_font(run, size=10.5, color="1F3A5F")
    body = paragraph.add_run(text)
    set_run_font(body, size=10.5, color="222222")
    document.add_paragraph()


def build_report() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Informe tecnico: Real-Time Finger Detection")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Proyecto Python con OpenCV y MediaPipe para contar dedos levantados en tiempo real"
    )
    meta = doc.add_paragraph()
    meta.add_run(f"Fecha: {date.today().strftime('%d/%m/%Y')} | Entorno: Windows | Proyecto: Real-Time Finger Detection")

    add_callout(
        doc,
        "Resumen",
        "La aplicacion captura video de la camara, detecta una o dos manos con MediaPipe, calcula los dedos extendidos a partir de los 21 landmarks y muestra el resultado sobre el video en vivo.",
    )

    doc.add_heading("1. Objetivo del proyecto", level=1)
    add_paragraph(
        doc,
        "El objetivo fue desarrollar una aplicacion de vision por computadora capaz de contar, en tiempo real, cuantos dedos estan levantados en una mano usando la camara del dispositivo. Durante la implementacion se amplio el soporte para detectar hasta dos manos.",
    )
    add_bullets(
        doc,
        [
            "Captura de video en vivo mediante OpenCV.",
            "Deteccion de manos con MediaPipe Hands.",
            "Conteo de dedos extendidos de 0 a 5 por mano.",
            "Overlay con landmarks, FPS y numero de dedos detectados.",
            "Suavizado temporal para reducir parpadeos en el conteo.",
        ]
    )

    doc.add_heading("2. Tecnologias utilizadas", level=1)
    add_table(
        doc,
        ["Tecnologia", "Uso en el proyecto"],
        [
            ["Python", "Lenguaje principal de la aplicacion y scripts de soporte."],
            ["OpenCV", "Captura de webcam, transformacion de frames y visualizacion en ventana."],
            ["MediaPipe", "Deteccion de mano y extraccion de 21 landmarks."],
            ["NumPy", "Dependencia numerica usada por OpenCV y MediaPipe."],
            ["python-docx", "Generacion de este informe tecnico en formato Word."],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("3. Estructura del proyecto", level=1)
    add_table(
        doc,
        ["Archivo", "Responsabilidad"],
        [
            ["main.py", "Punto de entrada. Gestiona argumentos, camara, ventana, overlay y controles."],
            ["hand_detector.py", "Encapsula MediaPipe, logica de dedos, dibujo de landmarks y soporte para recursos."],
            ["utils.py", "Incluye FPSCounter y StableCounter para metricas y estabilizacion."],
            ["requirements.txt", "Fija versiones compatibles de OpenCV, MediaPipe y NumPy."],
            ["README.md", "Instrucciones de instalacion, ejecucion y controles."],
        ],
        [1.65, 4.85],
    )

    doc.add_heading("4. Funcionamiento general", level=1)
    add_numbered(
        doc,
        [
            "Se abre la camara seleccionada con OpenCV y se configura resolucion y FPS objetivo.",
            "Cada frame se invierte horizontalmente por defecto para una experiencia tipo espejo.",
            "El frame BGR se convierte a RGB y se procesa con MediaPipe Hands.",
            "Si hay manos detectadas, se extraen landmarks y handedness.",
            "Para indice, medio, anular y menique se compara la coordenada Y de punta, PIP y MCP.",
            "Para el pulgar se compara la apertura lateral teniendo en cuenta si la mano es Left o Right.",
            "El conteo bruto se suaviza con una ventana temporal corta y se muestra en pantalla.",
        ]
    )

    doc.add_heading("5. Logica de conteo de dedos", level=1)
    add_paragraph(
        doc,
        "MediaPipe devuelve 21 puntos normalizados de la mano. La regla principal interpreta un dedo como levantado cuando la punta esta por encima de sus articulaciones intermedias. Esta aproximacion funciona mejor con la palma visible y los dedos orientados hacia arriba.",
    )
    add_table(
        doc,
        ["Dedo", "Landmark punta", "Comparacion principal"],
        [
            ["Pulgar", "4", "Apertura lateral frente a landmarks 3 y 2, ajustada por mano izquierda/derecha."],
            ["Indice", "8", "tip.y < pip.y < mcp.y"],
            ["Medio", "12", "tip.y < pip.y < mcp.y"],
            ["Anular", "16", "tip.y < pip.y < mcp.y"],
            ["Menique", "20", "tip.y < pip.y < mcp.y"],
        ],
        [1.3, 1.5, 3.7],
    )
    add_callout(
        doc,
        "Estabilidad",
        "El conteo final mostrado usa la moda de los ultimos frames. Esto evita saltos rapidos cuando un dedo queda cerca del limite entre levantado y doblado.",
    )

    doc.add_heading("6. Instalacion y ejecucion", level=1)
    add_paragraph(doc, "El proyecto quedo preparado con un entorno virtual local en la carpeta .venv.")
    add_code_block(
        doc,
        [
            ".\\.venv\\Scripts\\Activate.ps1",
            "pip install -r requirements.txt",
            "python main.py",
        ],
    )
    add_paragraph(doc, "Ejecucion directa sin activar el entorno:")
    add_code_block(doc, [".\\.venv\\Scripts\\python.exe main.py"])
    add_paragraph(doc, "Opciones utiles:")
    add_code_block(
        doc,
        [
            "python main.py --hands 1",
            "python main.py --hands 2",
            "python main.py --dominant-hand Right",
            "python main.py --thumb-mode flip",
        ],
    )

    doc.add_heading("7. Controles de usuario", level=1)
    add_table(
        doc,
        ["Tecla", "Accion"],
        [
            ["Q / ESC", "Cerrar la aplicacion."],
            ["R", "Reiniciar la ventana de estabilidad del conteo."],
            ["H", "Invertir la regla del pulgar si la camara lo cuenta al reves."],
        ],
        [1.4, 5.1],
    )

    doc.add_heading("8. Incidencias encontradas y resolucion", level=1)
    add_table(
        doc,
        ["Incidencia", "Causa", "Solucion aplicada"],
        [
            [
                "mediapipe no exponia mp.solutions",
                "La version mas reciente instalada no mantenia la API clasica esperada.",
                "Se fijo mediapipe==0.10.21 en requirements.txt.",
            ],
            [
                "Conflicto NumPy/OpenCV",
                "opencv-python moderno solicitaba NumPy 2 mientras MediaPipe clasico requiere NumPy 1.x.",
                "Se uso opencv-contrib-python==4.11.0.86 y numpy>=1.26,<2.",
            ],
            [
                "Fallo de recursos por ruta con acento",
                "MediaPipe en Windows fallo cargando grafos desde la antigua ruta DeteccionDedos con caracter acentuado.",
                "Se redirigieron los recursos de MediaPipe a una carpeta temporal ASCII antes de iniciar Hands.",
            ],
        ],
        [1.95, 2.05, 2.5],
    )

    doc.add_heading("9. Pruebas realizadas", level=1)
    add_table(
        doc,
        ["Prueba", "Resultado"],
        [
            ["Compilacion de modulos Python", "Correcta: main.py, hand_detector.py y utils.py compilan sin errores."],
            ["Import de dependencias", "Correcto: OpenCV 4.11.0, MediaPipe 0.10.21 y NumPy 1.26.4."],
            ["Inicializacion del detector", "Correcta: HandDetector se crea y se cierra sin excepciones."],
            ["Captura de camara", "Correcta: la webcam se abre y entrega frames de 480x640."],
            ["Prueba funcional del usuario", "Confirmada: la aplicacion funciona correctamente y detecta manos en vivo."],
        ],
        [2.45, 4.05],
    )

    doc.add_heading("10. Limitaciones y mejoras futuras", level=1)
    add_bullets(
        doc,
        [
            "La logica funciona mejor con palma visible, buena iluminacion y dedos orientados hacia arriba.",
            "El pulgar depende de la clasificacion Left/Right y puede requerir la tecla H segun la camara.",
            "Se podria anadir calibracion guiada para adaptar la postura inicial de cada usuario.",
            "Se podria guardar historico de FPS y precision estimada para sesiones de prueba.",
            "Se podria empaquetar como ejecutable para equipos sin entorno Python preparado.",
        ]
    )

    doc.add_heading("11. Conclusiones", level=1)
    add_paragraph(
        doc,
        "El proyecto cumple los requisitos planteados: usa camara en tiempo real, detecta manos, identifica dedos levantados, muestra el conteo sobre el video y mantiene una estructura clara por modulos. Tras los ajustes de dependencias y compatibilidad de rutas en Windows, la aplicacion quedo funcional y preparada para ampliaciones.",
    )

    doc.add_page_break()
    doc.add_heading("Anexo A. Dependencias fijadas", level=1)
    add_code_block(
        doc,
        [
            "opencv-contrib-python==4.11.0.86",
            "mediapipe==0.10.21",
            "numpy>=1.26.0,<2",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
